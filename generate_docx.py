#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成系统架构设计说明书 .docx 文件
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn

import os

doc = Document()

# ============================================================
# 样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading

def add_para(doc, text, bold=False, align=None, font_size=None, first_line_indent=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    return table

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_para(doc, '系统架构设计说明书', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=26)
doc.add_paragraph()
add_para(doc, 'Software Architecture Design Document', align=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)
doc.add_paragraph()
doc.add_paragraph()
add_para(doc, '项目名称：电影协同过滤推荐系统', align=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)
add_para(doc, '项目版本：v1.0', align=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)
add_para(doc, '编制日期：2026-06-26', align=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)

doc.add_page_break()

# ============================================================
# 文档修订记录
# ============================================================
add_heading_styled(doc, '文档修订记录', level=1)
add_table(doc, 
    ['版本号', '修订日期', '修订内容', '修订人'],
    [['1.0', '2026-06-26', '初始版本', '项目组']]
)

doc.add_page_break()

# ============================================================
# 目录占位
# ============================================================
add_heading_styled(doc, '目录', level=1)
add_para(doc, '（请在 Word 中右键此处 → 更新域，生成完整目录）', font_size=10)
doc.add_page_break()

# ============================================================
# 1. 概述
# ============================================================
add_heading_styled(doc, '1. 概述', level=1)

add_heading_styled(doc, '1.1 背景', level=2)
add_para(doc, '随着互联网视频平台的快速发展，用户面对海量电影内容时常常难以抉择。推荐系统作为解决信息过载问题的有效手段，能够帮助用户发现感兴趣的影片，提升用户体验和平台粘性。', first_line_indent=2)
add_para(doc, '本系统基于 MovieLens 100K 公开数据集，实现 Item-Based（基于物品）与 User-Based（基于用户）两种经典协同过滤推荐算法，并构建完整的前后端交互系统，为用户提供个性化的电影推荐服务。', first_line_indent=2)

add_heading_styled(doc, '1.2 范围', level=2)
add_para(doc, '本架构设计说明书覆盖以下系统范围：', first_line_indent=2)
items = [
    '数据概览模块：评分分布、用户/电影分布、类型分布、矩阵稀疏度展示',
    '相似电影模块：基于物品协同过滤的相似电影推荐',
    '个性化推荐模块：输入用户 ID，配置 K 值，生成 TopN 推荐列表',
    '模型评估模块：RMSE 计算、命中率评估、方法对比、相似度矩阵热力图展示',
    '数据管理模块：电影搜索、用户评分历史查看、热门电影排行'
]
for item in items:
    add_para(doc, f'• {item}', first_line_indent=2)

add_heading_styled(doc, '1.3 术语定义', level=2)
add_table(doc,
    ['术语', '英文', '说明'],
    [
        ['Item-Based CF', 'Item-Based Collaborative Filtering', '基于物品的协同过滤，通过计算物品间相似度进行推荐'],
        ['User-Based CF', 'User-Based Collaborative Filtering', '基于用户的协同过滤，通过计算用户间相似度进行推荐'],
        ['余弦相似度', 'Cosine Similarity', '衡量两个向量在方向上的一致性程度'],
        ['RMSE', 'Root Mean Square Error', '均方根误差，衡量预测评分与实际评分的偏差'],
        ['Hit Rate', 'Hit Rate', '命中率，推荐列表中用户实际喜欢的物品占比'],
        ['TopN', 'Top-N Recommendation', '为用户推荐评分最高的 N 个物品'],
        ['稀疏度', 'Sparsity', '评分矩阵中未评分项占总项的比例'],
        ['RESTful', 'Representational State Transfer', '使用 HTTP 方法操作资源的 API 设计规范'],
        ['SPA', 'Single Page Application', '单页面应用，前端路由由 JavaScript 控制']
    ]
)

add_heading_styled(doc, '1.4 参考资料', level=2)
refs = ['MovieLens 100K Dataset - GroupLens Research',
        'Scikit-learn 文档 - Cosine Similarity 实现参考',
        'Vue 3 官方文档 - 前端框架指南',
        'FastAPI 官方文档 - 后端框架参考',
        '《推荐系统实践》 - 项亮 编著']
for r in refs:
    add_para(doc, f'• {r}')

add_heading_styled(doc, '1.5 总体设计', level=2)
add_para(doc, '本系统采用前后端分离的架构设计模式。后端使用 FastAPI 构建 RESTful API 服务，前端使用 Vue 3 Composition API + Vite + Chart.js 构建用户交互界面。', first_line_indent=2)
add_para(doc, '核心推荐引擎基于 Pandas + NumPy 完成数据处理与矩阵运算，利用 Scikit-learn 提供的余弦相似度计算库实现相似度度量。数据源采用 GroupLens 提供的 MovieLens 100K 标准数据集。', first_line_indent=2)
add_para(doc, '系统总体架构分为三层：', first_line_indent=2)
layers = [
    '展示层（前端）：Vue 3 + Vue Router + Axios + Chart.js，负责数据可视化与用户交互',
    '业务逻辑层（后端 API）：FastAPI + Uvicorn，负责 API 路由、请求处理与数据聚合',
    '算法引擎层（推荐算法）：Pandas + NumPy + Scikit-learn，负责数据预处理、相似度计算、评分预测与模型评估'
]
for l in layers:
    add_para(doc, f'• {l}')

add_heading_styled(doc, '1.6 架构描述视图', level=2)
add_para(doc, '系统整体采用前后端分离的 MVC 架构模式：', first_line_indent=2)
add_para(doc, '• Model（模型）：MovieLens 数据集（CSV 文件）及内存中的数据对象（DataFrame、矩阵）', first_line_indent=2)
add_para(doc, '• View（视图）：Vue 3 前端组件，包含 5 个核心视图页面', first_line_indent=2)
add_para(doc, '• Controller（控制器）：FastAPI 路由处理函数，协调数据流与业务逻辑', first_line_indent=2)
add_para(doc, '前端渲染由 Vue 运行时完成，后端数据通过 API 异步获取，推荐算法结果在后端计算完成后以 JSON 格式返回前端展示。', first_line_indent=2)

# ============================================================
# 2. 系统架构策略
# ============================================================
add_heading_styled(doc, '2. 系统架构策略', level=1)

add_heading_styled(doc, '2.1 用户交互原则', level=2)
add_para(doc, '系统采用清晰简洁的浅新拟态（Soft Neumorphism）设计风格，以低饱和薄荷绿为主色，配合大面积留白和超大圆角卡片，提供友好的用户体验。', first_line_indent=2)
principles = [
    '实时反馈：用户操作后立即给出视觉反馈（加载态、成功/错误提示）',
    '渐进式加载：大数据量请求时展示进度状态，避免界面卡顿',
    '可视化优先：评分分布、相似度矩阵等数据以图表形式直观呈现',
    '响应式适配：支持不同屏幕尺寸的正常浏览'
]
for p in principles:
    add_para(doc, f'• {p}')

add_heading_styled(doc, '2.2 系统稳定性保证', level=2)
stability = [
    '异步 I/O：FastAPI 基于 Starlette，支持异步请求处理，防止 I/O 阻塞',
    '异常处理：全局异常捕获机制，统一返回错误格式',
    '数据校验：使用 Pydantic 模型对请求参数进行类型校验',
    '热重载开发：开发模式下支持代码修改自动重启',
    'CORS 配置：允许跨域请求，保障前后端分离开发'
]
for s in stability:
    add_para(doc, f'• {s}')

add_heading_styled(doc, '2.3 扩展性设计', level=2)
add_para(doc, '系统采用前后端分离的模块化架构设计：', first_line_indent=2)
add_para(doc, '前端扩展：组件化设计，各视图独立，可单独增加/修改推荐页面；Vue Router 支持动态路由；Chart.js 支持多种图表类型。', first_line_indent=2)
add_para(doc, '后端扩展：模块化路由组织，新增算法可独立注册路由；支持添加更多推荐算法（如 SVD、NCF 等）；数据结构支持更多数据源。', first_line_indent=2)
add_para(doc, '算法扩展：相似度度量可替换（余弦相似度、皮尔逊相关系数、Jaccard 相似度等）；预测公式可配置化；评估指标可扩展。', first_line_indent=2)

add_heading_styled(doc, '2.4 技术选型策略', level=2)
add_table(doc,
    ['层级', '技术', '选型理由'],
    [
        ['前端框架', 'Vue 3 (Composition API)', '轻量、高性能、组合式 API 便于逻辑复用'],
        ['构建工具', 'Vite 5', '快速冷启动、热更新、按需编译'],
        ['HTTP 客户端', 'Axios', '支持 Promise、拦截器、请求/响应转换'],
        ['图表库', 'Chart.js', '轻量、灵活、支持多种图表类型'],
        ['后端框架', 'FastAPI', '高性能异步框架、自动生成 API 文档'],
        ['数据处理', 'Pandas + NumPy', '高效的矩阵运算与数据操作能力'],
        ['机器学习', 'Scikit-learn', '成熟的余弦相似度计算、评估指标工具'],
        ['数据源', 'MovieLens 100K', '学术界标准基准数据集']
    ]
)

add_heading_styled(doc, '2.5 安全设计', level=2)
security = [
    'CORS 中间件：配置允许的前端来源域名，防止跨站请求伪造',
    '输入验证：Pydantic 模型校验用户输入，防止参数注入',
    '路径安全：文件路径使用安全拼接，防止目录遍历攻击',
    '数据校验：用户 ID、电影 ID 等参数格式校验',
    '错误信息控制：返回给前端的信息不包含堆栈细节，防止信息泄露'
]
for s in security:
    add_para(doc, f'• {s}')

add_heading_styled(doc, '2.6 开发环境与部署', level=2)
add_table(doc,
    ['组件', '技术', '版本'],
    [
        ['后端', 'FastAPI + Uvicorn', 'FastAPI 0.99+ / Uvicorn 0.22+'],
        ['前端', 'Vue 3 + Vite', 'Vue 3.3+ / Vite 5+'],
        ['Python', 'CPython', '3.8+'],
        ['Node.js', 'V8 引擎', '16+'],
        ['推荐算法依赖', 'Pandas / NumPy / Scikit-learn', '1.5+ / 1.24+ / 1.2+']
    ]
)
add_para(doc, '后端运行端口：8000；前端运行端口：5173', font_size=10)

# ============================================================
# 3. 关键数据结构设计
# ============================================================
add_heading_styled(doc, '3. 关键数据结构设计', level=1)

add_heading_styled(doc, '3.1 电影数据实体', level=2)
add_para(doc, '电影数据来源于 u.item 文件：', first_line_indent=2)
add_table(doc,
    ['字段名', '类型', '说明'],
    [
        ['movie_id', 'Integer', '电影唯一标识（1-1682）'],
        ['title', 'String', '电影标题'],
        ['release_date', 'Date', '上映日期'],
        ['video_release_date', 'Date', '录像带发行日期'],
        ['imdb_url', 'String', 'IMDb 链接地址'],
        ['genres[]', 'List[Integer]', '电影类型标签向量（0/1，共 19 种类型）']
    ]
)
add_para(doc, '类型标签：unknown、Action、Adventure、Animation、Children\'s、Comedy、Crime、Documentary、Drama、Fantasy、Film-Noir、Horror、Musical、Mystery、Romance、Sci-Fi、Thriller、War、Western')

add_heading_styled(doc, '3.2 用户数据实体', level=2)
add_table(doc,
    ['字段名', '类型', '说明'],
    [
        ['user_id', 'Integer', '用户唯一标识（1-943）'],
        ['age', 'Integer', '用户年龄'],
        ['gender', 'String', '用户性别（M/F）'],
        ['occupation', 'String', '职业类别'],
        ['zip_code', 'String', '邮政编码']
    ]
)

add_heading_styled(doc, '3.3 评分数据实体', level=2)
add_table(doc,
    ['字段名', '类型', '说明'],
    [
        ['user_id', 'Integer', '用户 ID'],
        ['item_id', 'Integer', '电影 ID'],
        ['rating', 'Integer', '评分（1-5 分）'],
        ['timestamp', 'Integer', '时间戳（Unix 时间）']
    ]
)

add_heading_styled(doc, '3.4 相似度矩阵实体', level=2)
add_table(doc,
    ['字段', '类型', '说明'],
    [
        ['matrix_type', 'String', '相似度类型（item_similarity / user_similarity）'],
        ['shape', 'Tuple[int,int]', '矩阵维度'],
        ['data', 'ndarray', '相似度矩阵数据（浮点数）'],
        ['index_map', 'Dict[int,int]', 'ID 到矩阵索引的映射'],
        ['method', 'String', '计算方法（cosine）']
    ]
)

add_heading_styled(doc, '3.5 推荐结果实体', level=2)
add_table(doc,
    ['字段名', '类型', '说明'],
    [
        ['user_id', 'Integer', '被推荐用户 ID'],
        ['recommendations', 'List[Dict]', '推荐电影列表 [{movie_id, title, predicted_rating}]'],
        ['method', 'String', '推荐方法（item_based / user_based）'],
        ['k', 'Integer', '邻居数量'],
        ['computation_time', 'Float', '计算耗时（秒）']
    ]
)

# ============================================================
# 4. 数据存储设计
# ============================================================
add_heading_styled(doc, '4. 数据存储设计', level=1)

add_heading_styled(doc, '4.1 MovieLens 数据集结构', level=2)
add_para(doc, '系统使用 MovieLens 100K 数据集，以 CSV 格式存储在 ml-100k/ 目录下：', first_line_indent=2)
add_table(doc,
    ['文件名', '记录数', '说明'],
    [
        ['u.data', '100,000 条', '评分数据（user_id, item_id, rating, timestamp）'],
        ['u.item', '1,682 条', '电影信息（ID、标题、上映日期、类型标签等）'],
        ['u.user', '943 条', '用户信息（ID、年龄、性别、职业等）'],
        ['u.genre', '19 条', '电影类型列表'],
        ['u.occupation', '21 条', '用户职业列表']
    ]
)
add_para(doc, '数据统计：用户数 943，电影数 1,682，评分记录数 100,000，评分范围 1-5，稀疏度约 93.7%。')

add_heading_styled(doc, '4.2 内存数据结构', level=2)
add_para(doc, '推荐算法运行时主要使用以下内存数据结构：', first_line_indent=2)
add_para(doc, '评分矩阵（Rating Matrix）：Pandas DataFrame（943×1682），行索引为 user_id，列索引为 movie_id，值 NaN 表示未评分。', first_line_indent=2)
add_para(doc, '相似度矩阵（Similarity Matrix）：Item-Based 为 1682×1682 对称矩阵，User-Based 为 943×943 对称矩阵，值范围 [0, 1]，使用余弦相似度计算。', first_line_indent=2)

add_heading_styled(doc, '4.3 缓存设计', level=2)
add_para(doc, '系统采用内存缓存策略提升性能，缓存策略采用延迟计算（Lazy Computation）模式：', first_line_indent=2)
add_table(doc,
    ['缓存对象', '更新时机', '说明'],
    [
        ['评分矩阵', '服务启动时加载', '从 u.data 读取，构建 user-item 矩阵'],
        ['电影信息表', '服务启动时加载', '从 u.item 读取，用于展示电影详情'],
        ['Item 相似度矩阵', '首次请求时计算并缓存', '计算量大，仅在服务运行期间缓存'],
        ['User 相似度矩阵', '首次请求时计算并缓存', '计算量大，仅在服务运行期间缓存']
    ]
)

add_heading_styled(doc, '4.4 数据预处理流程', level=2)
add_para(doc, '数据预处理流程：原始 CSV 文件 → Pandas 读取 → 数据清洗 → 构建评分矩阵 → 导出统计数据。', first_line_indent=2)
steps = [
    '数据加载：使用 Pandas 读取 u.data、u.item、u.user 文件',
    '缺失值处理：检查并处理缺失数据',
    '类型转换：将评分转为整数，时间戳转为 datetime',
    '用户-物品矩阵构建：创建 user×movie 的评分矩阵（NaN 填充未评分）',
    '数据统计：计算用户数、电影数、评分总数、稀疏度、评分分布等',
    '缓存就绪：将预处理结果存入内存，等待 API 调用'
]
for s in steps:
    add_para(doc, f'• {s}')

# ============================================================
# 5. 接口设计
# ============================================================
add_heading_styled(doc, '5. 接口设计', level=1)

add_heading_styled(doc, '5.1 API 接口列表', level=2)
add_para(doc, '系统后端提供以下 RESTful API：', first_line_indent=2)
add_table(doc,
    ['方法', '路径', '说明', '参数'],
    [
        ['GET', '/api/stats', '数据统计', '-'],
        ['GET', '/api/movies', '电影列表（支持搜索/分页）', 'q, page, page_size'],
        ['GET', '/api/movies/{id}', '电影详情', 'id'],
        ['GET', '/api/movies/{id}/similar', '相似电影推荐', 'id, k'],
        ['GET', '/api/recommend/{user_id}', '个性化推荐', 'user_id, k, n'],
        ['GET', '/api/users/{id}/ratings', '用户评分历史', 'id'],
        ['GET', '/api/ratings/distribution', '评分分布', '-'],
        ['GET', '/api/genres/distribution', '电影类型分布', '-'],
        ['GET', '/api/popular-movies', '热门电影排行', 'n'],
        ['GET', '/api/similarity-matrix', '相似度矩阵数据', 'type'],
        ['POST', '/api/evaluate/rmse', 'RMSE 评估', 'method, k, test_size'],
        ['POST', '/api/evaluate/hit-rate', '命中率评估', 'method, k, n, test_size']
    ]
)

add_heading_styled(doc, '5.2 前端与后端集成方式', level=2)
add_para(doc, '前端使用 Axios 统一发起 HTTP 请求到后端 API：', first_line_indent=2)
add_para(doc, '请求流程：Vue 组件 → API 封装层 (api/index.js) → Axios GET/POST → FastAPI 路由 → Pandas/NumPy 处理 → JSON 响应 → Axios 解析 → 组件状态更新 → DOM 渲染。', first_line_indent=2)
add_para(doc, '通用响应格式：{ "code": 200, "message": "success", "data": { ... } }', first_line_indent=2)
add_para(doc, '开发代理配置：Vite 开发服务器将 /api 请求代理到 http://localhost:8000。', first_line_indent=2)

# ============================================================
# 6. 部署视图
# ============================================================
add_heading_styled(doc, '6. 部署视图', level=1)

add_heading_styled(doc, '6.1 系统架构总览', level=2)
add_para(doc, '系统整体架构分为三层：', first_line_indent=2)
add_para(doc, '第一层为用户浏览器中的 Vue 3 前端 SPA（单页面应用），包含首页、数据概览、相似电影、个性化推荐、模型评估 5 个核心视图页面。', first_line_indent=2)
add_para(doc, '第二层为 FastAPI 后端服务，负责 API 路由处理、数据聚合、异常处理和跨域配置。', first_line_indent=2)
add_para(doc, '第三层为推荐算法引擎，基于 Pandas/NumPy/Scikit-learn 实现 Item-Based CF 和 User-Based CF 两种协同过滤算法，并支持 RMSE 和 Hit Rate 评估。', first_line_indent=2)
add_para(doc, '数据源为 MovieLens 100K 数据集，以 CSV 文件形式存储在 ml-100k/ 目录下。', first_line_indent=2)

add_heading_styled(doc, '6.2 前端架构', level=2)
add_para(doc, '前端使用 Vue 3 Composition API 构建，路由由 Vue Router 管理：', first_line_indent=2)
add_table(doc,
    ['视图', '路由路径', '核心功能'],
    [
        ['首页 (Home)', '/', '系统概览、快速入口'],
        ['数据概览 (DataOverview)', '/data-overview', '评分分布、用户/电影分布、类型分布、稀疏度'],
        ['相似电影 (SimilarMovies)', '/similar-movies', '输入电影 ID，查看相似电影列表'],
        ['个性化推荐 (Recommendations)', '/recommendations', '输入用户 ID，配置 K/N 值，生成推荐'],
        ['模型评估 (Evaluation)', '/evaluation', 'RMSE 评估、命中率评估、方法对比、热力图']
    ]
)

add_heading_styled(doc, '6.3 后端架构', level=2)
add_para(doc, '后端使用 FastAPI 构建，采用模块化路由组织。核心 API 涵盖数据统计、电影查询、推荐生成和模型评估四大类，共 12 个接口。', first_line_indent=2)
add_para(doc, '后端在启动时加载数据集（u.data、u.item、u.user）到内存，以 Pandas DataFrame 形式存储。相似度矩阵采用延迟加载策略，仅在首次请求时计算并缓存。', first_line_indent=2)

add_heading_styled(doc, '6.4 算法引擎架构', level=2)
add_para(doc, '推荐算法引擎为核心计算模块，包含以下功能：', first_line_indent=2)
add_para(doc, '数据预处理：构建 user-item 评分矩阵（943×1682），评分归一化去均值，训练/测试集划分。', first_line_indent=2)
add_para(doc, '相似度计算：计算物品间两两余弦相似度（1682×1682 矩阵），或用户间两两余弦相似度（943×943 矩阵），取 Top-K 邻居。', first_line_indent=2)
add_para(doc, '评分预测：Item-Based 使用加权平均预测，User-Based 使用去均值加权预测。', first_line_indent=2)
add_para(doc, 'TopN 推荐：预测所有未评分电影评分，取 Top-N 作为推荐结果。', first_line_indent=2)
add_para(doc, '模型评估：支持 RMSE 和 Hit Rate 两类评估指标，支持两种算法方法的横向对比。', first_line_indent=2)

add_heading_styled(doc, '6.5 推荐算法流程', level=2)
add_para(doc, '推荐流程如下：', first_line_indent=2)
steps = [
    '用户输入 user_id，选择 K 值和 N 值',
    '后端加载评分矩阵（如未缓存则从文件读取）',
    '获取用户已评分的电影列表',
    '计算（或加载缓存）相似度矩阵',
    '对每部未评分电影，找到用户已评分电影中与当前电影最相似的 K 个邻居',
    '基于邻居的评分加权预测用户对当前电影的评分',
    '按预测评分降序排列，取前 N 部电影',
    '补充电影详情（标题、类型等），返回 JSON 结果'
]
for i, s in enumerate(steps, 1):
    add_para(doc, f'{i}. {s}')

add_para(doc, '性能基准：数据加载约 200ms（首次），相似度计算约 500ms-2s（首次），单次推荐预测约 50-200ms（基于缓存），模型评估约 1-5s。', first_line_indent=2)

# ============================================================
# 页脚信息
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
add_para(doc, '— 文档结束 —', align=WD_ALIGN_PARAGRAPH.CENTER, font_size=10)
add_para(doc, '文档版本：v1.0 | 最后更新：2026-06-26 | 项目名称：电影协同过滤推荐系统', align=WD_ALIGN_PARAGRAPH.CENTER, font_size=9)

# ============================================================
# 保存
# ============================================================
output_path = r'D:\vscode project\group project\系统架构设计说明书.docx'
doc.save(output_path)
print(f'文档已成功生成: {output_path}')
