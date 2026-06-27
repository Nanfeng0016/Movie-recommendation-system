"""
Markdown → Docx 转换脚本
将系统构架设计说明书.md 转换为 Word 文档
支持标题层级、表格、代码块、列表等元素
"""

import re
import os

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("正在安装 python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn


def add_table_from_md(doc, table_lines):
    """将 Markdown 表格转换为 Word 表格"""
    rows = []
    for line in table_lines:
        if line.strip().startswith('|---') or line.strip().startswith('|:---'):
            continue
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line.strip('|').split('|')]
            rows.append(cells)

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = cell_text
            # 表头加粗
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()  # 表后空行


def add_code_block(doc, code_text, language=""):
    """添加代码块（灰色底纹、等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    # 为代码块添加底纹
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F0F0')
    shading.set(qn('w:val'), 'clear')
    p.paragraph_format.element.get_or_add_pPr().append(shading)

    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def convert_md_to_docx(md_path, docx_path):
    """主转换函数"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []
    in_list = False
    list_type = None

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # 代码块处理
        if stripped.startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(stripped)
            i += 1
            continue

        # 表格处理
        if stripped.startswith('|') and stripped.endswith('|'):
            table_buffer.append(stripped)
            in_table = True
            i += 1
            continue
        else:
            if in_table:
                add_table_from_md(doc, table_buffer)
                table_buffer = []
                in_table = False

        # 标题处理
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
            i += 1
            continue
        if stripped.startswith('##### '):
            doc.add_heading(stripped[6:], level=5)
            i += 1
            continue

        # 分割线
        if stripped.startswith('---') or stripped.startswith('***'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            from docx.oxml import OxmlElement
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '999999')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # 列表项
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            # 处理粗体
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        if re.match(r'^\d+\. ', stripped):
            text = re.sub(r'^\d+\. ', '', stripped)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # 普通段落
        if stripped:
            # 处理粗体
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            # 处理行内代码
            text = re.sub(r'`(.+?)`', r'\1', text)
            # 处理链接 [text](url)
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
            p = doc.add_paragraph(text)
            p.paragraph_format.line_spacing = Pt(20)
        else:
            # 空行
            p = doc.add_paragraph('')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

        i += 1

    # 处理末尾未闭合的表格
    if in_table:
        add_table_from_md(doc, table_buffer)

    # 保存
    doc.save(docx_path)
    print(f"✅ 转换完成！输出文件：{docx_path}")


if __name__ == '__main__':
    md_file = r"D:\vscode project\group project\系统构架设计说明书.md"
    docx_file = r"D:\vscode project\group project\系统构架设计说明书.docx"
    
    if not os.path.exists(md_file):
        print(f"❌ 未找到文件：{md_file}")
    else:
        convert_md_to_docx(md_file, docx_file)
